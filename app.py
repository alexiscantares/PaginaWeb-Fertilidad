import io
import json
import base64
from pathlib import Path

import streamlit as st
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.interpolate import griddata
from sklearn.impute import SimpleImputer
from PIL import Image
from groq import Groq
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =====================================================================
# RUTAS ABSOLUTAS DINÁMICAS
# =====================================================================
BASE_DIR = Path(__file__).resolve().parent

# =====================================================================
# CONFIGURACIÓN DE PÁGINA Y ESTILOS
# =====================================================================
st.set_page_config(page_title="Fertilidad Geoquímica NI 43-101", layout="wide", page_icon="🌋")
sns.set_theme(style="whitegrid")

st.title("Clasificador de Fertilidad Magmática & Generador NI 43-101")
st.markdown("Plataforma de inferencia multivariada, interpretación metalogénica y reporte técnico libre.")
st.markdown("---")

# =====================================================================
# CARGA DE MODELOS (CACHÉ OPTIMIZADA)
# =====================================================================
@st.cache_resource
def cargar_modelos():
    ruta_scaler = BASE_DIR / 'scaler_geoquimico.pkl'
    ruta_modelo = BASE_DIR / 'modelo_fertilidad_rf.pkl'
    
    scaler = joblib.load(ruta_scaler)
    modelo = joblib.load(ruta_modelo)
    return scaler, modelo

try:
    scaler, modelo_rf = cargar_modelos()
except Exception as e:
    st.error(f"❌ Error al cargar los modelos (.pkl): {e}")
    st.stop()

elementos_requeridos = [
    'AU', 'AG', 'CU', 'PB', 'ZN', 'MO', 'NI', 'CO', 'CD', 'BI',
    'FE', 'MN', 'TE', 'BA', 'CR', 'V', 'SN', 'W', 'LA', 'AL',
    'MG', 'CA', 'NA', 'K', 'SR', 'Y', 'GA', 'LI', 'NB', 'SC',
    'TA', 'TI', 'ZR', 'AS', 'SB', 'HG', 'PT', 'PD'
]

# =====================================================================
# PANEL LATERAL (SIDEBAR)
# =====================================================================
st.sidebar.header("⚙️ Panel de Control")
st.sidebar.write("Sube la matriz analítica del laboratorio.")
archivo_subido = st.sidebar.file_uploader("Formato CSV o Excel", type=['csv', 'xlsx'])

st.sidebar.markdown("---")
st.sidebar.header("🎛️ Parámetros del Modelo")
umbral_corte = st.sidebar.slider("Umbral de Probabilidad (Corte Fértil)", 0.0, 1.0, 0.5, 0.05)
st.sidebar.markdown("---")

def fig_a_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)  # Libera memoria de Matplotlib
    return buf

# =====================================================================
# FUNCIÓN GENERADORA DE DOCUMENTO WORD CON GRÁFICOS MULTIVARIADOS
# =====================================================================
def generar_documento_word(texto_reporte, figuras_dict):
    doc = docx.Document()
    
    # Encabezado Principal
    doc.add_heading("INFORME TÉCNICO DE EXPLORACIÓN GEOQUÍMICA (NI 43-101)", level=0)
    p_sub = doc.add_paragraph()
    p_sub.add_run("Evaluación de Fertilidad Magmática e Inferencia Multivariada por IA\n").bold = True

    def insertar_grafico(key_fig, titulo_fig, desc_fig):
        if key_fig in figuras_dict and figuras_dict[key_fig] is not None:
            buf = figuras_dict[key_fig]
            buf.seek(0)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(buf, width=Inches(5.5))
            
            p_label = doc.add_paragraph()
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_lbl = p_label.add_run(f"{titulo_fig}: {desc_fig}")
            run_lbl.italic = True
            run_lbl.font.size = Pt(9.5)
            p_label.paragraph_format.space_after = Pt(14)

    # Control de inserción única de imágenes
    insertadas = {
        "mapa": False,
        "sry": False,
        "spider": False,
        "afm": False,
        "cuk": False,
        "vti": False
    }

    lineas = texto_reporte.split('\n')
    
    for linea in lineas:
        linea_clean = linea.strip()
        if not linea_clean:
            continue

        es_titulo = any(sec in linea_clean.upper() for sec in [
            "1.", "2.", "3.", "4.", "5.", 
            "RESUMEN EJECUTIVO", "FIRMA ADAKÍTICA", 
            "EVALUACIÓN DE ELEMENTOS", "ALTERACIÓN GEOQUÍMICA", "CONCLUSIONES"
        ])

        if es_titulo and linea_clean.startswith("#"):
            doc.add_heading(linea_clean.replace("#", "").replace("**", "").strip(), level=2)
        elif es_titulo and not linea_clean.startswith("#") and len(linea_clean) < 80:
            doc.add_heading(linea_clean.replace("**", "").strip(), level=2)
        else:
            doc.add_paragraph(linea_clean.replace("**", ""))

        # Mapeo y control estricto de una sola inserción por tema
        if ("1." in linea_clean or "RESUMEN EJECUTIVO" in linea_clean.upper()) and not insertadas["mapa"]:
            insertar_grafico("mapa", "Figura 1. Modelo Digital Espacial de Anomalías", "Interpolación espacial de la probabilidad de fertilidad magmática e isolíneas de prospección.")
            insertadas["mapa"] = True
        
        elif ("2." in linea_clean or "FIRMA ADAKÍTICA" in linea_clean.upper()) and not insertadas["sry"]:
            insertar_grafico("sry", "Figura 2. Diagrama Sr/Y vs Y", "Relación de fraccionamiento de granate/anfíbol indicativo de la fertilidad magmática.")
            insertadas["sry"] = True
        
        elif ("3." in linea_clean or "EVALUACIÓN DE ELEMENTOS" in linea_clean.upper() or "METALOVECTORES" in linea_clean.upper()) and not insertadas["spider"]:
            insertar_grafico("spider", "Figura 3. Perfil Multi-elemental Normalizado", "Spider diagram comparativo entre la población fértil y el background estéril.")
            insertadas["spider"] = True
        
        elif ("4." in linea_clean or "ALTERACIÓN GEOQUÍMICA" in linea_clean.upper()) and not insertadas["afm"]:
            insertar_grafico("afm", "Figura 4. Diagrama Ternario AFM", "Proyección ternaria Na+K - Fe - Mg para series magmáticas.")
            insertar_grafico("cuk", "Figura 5. Relación Cu vs K", "Indicador de enriquecimiento potásico asociado a sulfuros de cobre.")
            insertar_grafico("vti", "Figura 6. Diagrama V vs Ti", "Condiciones de fugacidad de oxígeno (fO2) y fraccionamiento.")
            insertadas["afm"] = True
            insertadas["cuk"] = True
            insertadas["vti"] = True

    buffer_word = io.BytesIO()
    doc.save(buffer_word)
    buffer_word.seek(0)
    return buffer_word

# =====================================================================
# ÁREA PRINCIPAL DE PROCESAMIENTO
# =====================================================================
if 'modelo_ejecutado' not in st.session_state:
    st.session_state.modelo_ejecutado = False

if archivo_subido is not None:
    if archivo_subido.name.endswith('.csv'):
        df_input = pd.read_csv(archivo_subido)
    else:
        df_input = pd.read_excel(archivo_subido)
        
    st.sidebar.success(f"Archivo cargado: {len(df_input)} muestras.")
    
    if st.sidebar.button("🚀 Ejecutar Modelo Predictivo"):
        st.session_state.modelo_ejecutado = True
        
    if st.session_state.modelo_ejecutado:
        with st.spinner('Procesando datos y ejecutando análisis multivariado...'):
            try:
                # Verificar columnas faltantes
                columnas_faltantes = [elem for elem in elementos_requeridos if elem not in df_input.columns]
                if columnas_faltantes:
                    st.error(f"⚠️ El archivo subido no contiene los siguientes elementos requeridos: {', '.join(columnas_faltantes)}")
                    st.stop()

                # 1. Imputación e Inferencia
                datos_modelo = df_input[elementos_requeridos].copy()
                nans_detectados = datos_modelo.isna().sum().sum()
                
                if nans_detectados > 0:
                    imputer = SimpleImputer(strategy='median')
                    datos_modelo[elementos_requeridos] = imputer.fit_transform(datos_modelo)
                    datos_modelo = datos_modelo.fillna(0)
                    df_input[elementos_requeridos] = datos_modelo[elementos_requeridos]

                datos_log = np.log10(datos_modelo + 1e-5).fillna(0)
                datos_escalados = scaler.transform(datos_log)
                probabilidades = modelo_rf.predict_proba(datos_escalados)[:, 1]
                
                df_input['Prob_Fertilidad'] = probabilidades
                df_input['Clasificacion_IA'] = np.where(df_input['Prob_Fertilidad'] >= umbral_corte, 'Fértil', 'Estéril/Artefacto')
                
                if 'SR' in df_input.columns and 'Y' in df_input.columns:
                    df_input['Sr_Y'] = df_input['SR'] / (df_input['Y'] + 1e-5)

                st.session_state.figuras_informe = {}

                # =====================================================================
                # CREACIÓN DE PESTAÑAS (TABS)
                # =====================================================================
                tab1, tab2, tab3, tab4 = st.tabs(["📄 Resumen y Datos", "📉 Diagramas Geoquímicos", "🗺️ Mapa Espacial de Anomalías", "🤖 Asistente Groq NI 43-101"])
                
                # ----- PESTAÑA 1: RESUMEN Y DATOS -----
                with tab1:
                    st.subheader("📊 Resumen Analítico de la Campaña")
                    total_muestras = len(df_input)
                    muestras_fertiles = len(df_input[df_input['Clasificacion_IA'] == 'Fértil'])
                    porcentaje_anomalias = (muestras_fertiles / total_muestras) * 100 if total_muestras > 0 else 0
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Total Muestras Procesadas", total_muestras)
                    col2.metric("Blancos Fértiles Detectados", muestras_fertiles)
                    col3.metric("Tasa de Anomalía Geoquímica", f"{porcentaje_anomalias:.2f}%")
                    
                    st.markdown("---")
                    st.markdown("### 📋 Tabla de Inferencia Multivariada")
                    
                    def colorear_fertiles(val):
                        return 'background-color: #ffcccc' if val == 'Fértil' else ''
                    
                    st.dataframe(df_input.head(500).style.map(colorear_fertiles, subset=['Clasificacion_IA']), use_container_width=True)

                # ----- PESTAÑA 2: DIAGRAMAS COMPILADOS CON CAPTURA DE IMAGEN -----
                with tab2:
                    st.subheader("📉 Evaluaciones Geoquímicas Individuales")
                    cmap_prob = 'coolwarm'
                    
                    # Figura Sr/Y vs Y
                    st.markdown("### Figura: Diagrama de Fertilidad Magmática ($Sr/Y$ vs $Y$)")
                    fig1, ax1 = plt.subplots(figsize=(8, 4.5))
                    sns.scatterplot(data=df_input, x='Y', y='Sr_Y', hue='Prob_Fertilidad', palette=cmap_prob, s=30, alpha=0.8, ax=ax1)
                    ax1.set_xscale('log')
                    ax1.set_yscale('log')
                    ax1.axhline(20, color='red', linestyle='--', label='Corte Adakítico (Sr/Y = 20)')
                    ax1.set_xlabel('Yttrium (Y) [ppm]')
                    ax1.set_ylabel('Ratio Sr / Y')
                    ax1.legend(loc='best')
                    st.pyplot(fig1)
                    st.session_state.figuras_informe["sry"] = fig_a_bytes(fig1)

                    st.markdown("---")

                    # Figura Spider Diagram
                    st.markdown("### Figura: Perfil Multi-elemental Normalizado (Spider Diagram)")
                    fig2, ax2 = plt.subplots(figsize=(9, 4.5))
                    trace_elems = [e for e in ['LA', 'SR', 'Y', 'ZR', 'TI', 'V', 'SC', 'CU', 'ZN', 'PB'] if e in df_input.columns]
                    means_fert = df_input[df_input['Clasificacion_IA'] == 'Fértil'][trace_elems].mean()
                    means_inf = df_input[df_input['Clasificacion_IA'] != 'Fértil'][trace_elems].mean()
                    ax2.plot(trace_elems, np.log10(means_fert + 1e-5), marker='o', color='darkred', lw=2, label='Población Fértil')
                    ax2.plot(trace_elems, np.log10(means_inf + 1e-5), marker='s', color='darkblue', lw=2, label='Población Estéril')
                    ax2.set_ylabel('Concentración Log10 (ppm)')
                    ax2.legend()
                    st.pyplot(fig2)
                    st.session_state.figuras_informe["spider"] = fig_a_bytes(fig2)

                    st.markdown("---")

                    # Figura Ternario AFM
                    st.markdown("### Figura: Diagrama Ternario AFM ($Na+K$ - $Fe$ - $Mg$)")
                    fig3, ax3 = plt.subplots(figsize=(7, 5))
                    A, F, M = df_input['NA'] + df_input['K'], df_input['FE'], df_input['MG']
                    Total = A + F + M + 1e-5
                    X_tern = (M/Total) + (F/Total) / 2.0
                    Y_tern = (F/Total) * np.sqrt(3) / 2.0
                    sc = ax3.scatter(X_tern, Y_tern, c=df_input['Prob_Fertilidad'], cmap=cmap_prob, s=25, alpha=0.8)
                    ax3.plot([0, 1, 0.5, 0], [0, 0, np.sqrt(3)/2, 0], 'k-', lw=1.5)
                    ax3.text(-0.05, -0.05, 'A (Na+K)', ha='center')
                    ax3.text(1.05, -0.05, 'M (Mg)', ha='center')
                    ax3.text(0.5, np.sqrt(3)/2 + 0.05, 'F (Fe)', ha='center')
                    ax3.axis('off')
                    fig3.colorbar(sc, ax=ax3, label='Probabilidad IA')
                    st.pyplot(fig3)
                    st.session_state.figuras_informe["afm"] = fig_a_bytes(fig3)

                    st.markdown("---")

                    # Figura Cu vs K
                    st.markdown("### Figura: Relación Cu vs K (Indicador de Alteración Potásica)")
                    fig4, ax4 = plt.subplots(figsize=(8, 4.5))
                    sns.scatterplot(data=df_input, x='K', y='CU', hue='Prob_Fertilidad', palette=cmap_prob, s=30, alpha=0.8, ax=ax4)
                    ax4.set_xscale('log')
                    ax4.set_yscale('log')
                    ax4.set_xlabel('Potasio (K) [%]')
                    ax4.set_ylabel('Cobre (Cu) [ppm]')
                    st.pyplot(fig4)
                    st.session_state.figuras_informe["cuk"] = fig_a_bytes(fig4)

                    st.markdown("---")

                    # Figura V vs Ti
                    st.markdown("### Figura: Diagrama V vs Ti (Estado de Oxidación y Fraccionamiento)")
                    fig5, ax5 = plt.subplots(figsize=(8, 4.5))
                    sns.scatterplot(data=df_input, x='TI', y='V', hue='Prob_Fertilidad', palette=cmap_prob, s=30, alpha=0.8, ax=ax5)
                    ax5.set_xscale('log')
                    ax5.set_yscale('log')
                    ax5.set_xlabel('Titanio (Ti) [ppm]')
                    ax5.set_ylabel('Vanadio (V) [ppm]')
                    st.pyplot(fig5)
                    st.session_state.figuras_informe["vti"] = fig_a_bytes(fig5)

                # ----- PESTAÑA 3: MAPA GEOQUÍMICO ESPACIAL DE ANOMALÍAS -----
                with tab3:
                    st.subheader("🗺️ Mapa Geoquímico Espacial de Anomalías y Prospectividad")
                    if 'LONGITUD' in df_input.columns and 'LATITUD' in df_input.columns:
                        df_mapa = df_input.dropna(subset=['LATITUD', 'LONGITUD'])
                        if len(df_mapa) > 0:
                            X_coord = df_mapa['LONGITUD'].values
                            Y_coord = df_mapa['LATITUD'].values
                            Z_prob = df_mapa['Prob_Fertilidad'].values
                            
                            grid_x, grid_y = np.mgrid[X_coord.min():X_coord.max():200j, Y_coord.min():Y_coord.max():200j]
                            grid_z = griddata((X_coord, Y_coord), Z_prob, (grid_x, grid_y), method='linear')
                            
                            fig_map, ax_map = plt.subplots(figsize=(10, 7))
                            mapa_calor = ax_map.imshow(grid_z.T, extent=(X_coord.min(), X_coord.max(), Y_coord.min(), Y_coord.max()),
                                                       origin='lower', cmap='coolwarm', alpha=0.85)
                            
                            contornos = ax_map.contour(grid_x, grid_y, grid_z, levels=[0.5, 0.7, 0.9], colors=['yellow', 'orange', 'darkred'], linewidths=1.5, linestyles='--')
                            ax_map.clabel(contornos, inline=True, fontsize=9, fmt='P: %.1f')
                            
                            ax_map.scatter(X_coord, Y_coord, c='black', s=10, alpha=0.4, label='Estaciones de Muestreo')
                            
                            ax_map.set_title('Modelo Digital Espacial de Probabilidad de Fertilidad Magmática', fontsize=12, fontweight='bold')
                            ax_map.set_xlabel('Longitud (WGS84)')
                            ax_map.set_ylabel('Latitud (WGS84)')
                            cbar = fig_map.colorbar(mapa_calor, ax=ax_map, shrink=0.8)
                            cbar.set_label('Índice de Fertilidad (Probabilidad IA)')
                            ax_map.legend(loc='upper right')
                            
                            st.pyplot(fig_map)
                            st.session_state.figuras_informe["mapa"] = fig_a_bytes(fig_map)
                    else:
                        st.info("ℹ️ Para desplegar el mapa espacial, asegúrate de incluir las columnas 'LONGITUD' y 'LATITUD' en tu archivo.")

                # ----- PESTAÑA 4: ASISTENTE GROQ & GENERACIÓN NI 43-101 -----
                with tab4:
                    st.subheader("🤖 Generador de Informe Técnico Norma NI 43-101")
                    
                    if st.button("📄 Sintetizar e Interpretar Informe NI 43-101"):
                        groq_api_key = st.secrets.get("GROQ_API_KEY", "").strip()
                        
                        if not groq_api_key:
                            st.error("❌ Error de configuración del servidor: No se encontró 'GROQ_API_KEY' en los Secretos de Streamlit.")
                        else:
                            with st.spinner("Procesando síntesis con Llama 3.3 70B en Groq..."):
                                try:
                                    total_muestras_tab = len(df_input)
                                    df_fertil = df_input[df_input['Clasificacion_IA'] == 'Fértil']
                                    df_esteril = df_input[df_input['Clasificacion_IA'] != 'Fértil']
                                    muestras_fertiles_tab = len(df_fertil)
                                    porcentaje_anomalias_tab = (muestras_fertiles_tab / total_muestras_tab) * 100 if total_muestras_tab > 0 else 0
                                    
                                    trace_elems_g = [e for e in ['LA', 'SR', 'Y', 'ZR', 'TI', 'V', 'SC', 'CU', 'ZN', 'PB', 'K', 'FE', 'CR'] if e in df_input.columns]
                                    prom_fertil = df_fertil[trace_elems_g].mean().round(2).to_dict() if not df_fertil.empty else {}
                                    prom_esteril = df_esteril[trace_elems_g].mean().round(2).to_dict() if not df_esteril.empty else {}
                                    sry_mediano_fertil = round(df_fertil['Sr_Y'].median(), 2) if 'Sr_Y' in df_fertil.columns and not df_fertil.empty else "N/A"

                                    prompt_ni43101 = f"""
Actúa como una Persona Calificada (QP - Qualified Person) en exploración geoquímica bajo los estándares de la norma canadiense NI 43-101.

Redacta un informe técnico detallado enfocado en el ÍTEM 9 (EXPLORACIÓN) basado en los datos multivariados ingresados:

RESUMEN ANALÍTICO DE LA CAMPAÑA:
- Muestras Totales: {total_muestras_tab}
- Muestras Anómalas/Fértiles: {muestras_fertiles_tab} ({porcentaje_anomalias_tab:.2f}% de la población)
- Umbral de Probabilidad Aplicado: {umbral_corte}
- Mediana del Ratio Sr/Y en Zona Fértil: {sry_mediano_fertil}

PROMEDIOS ELEMENTALES (Fértil vs Estéril):
- Fértil: {prom_fertil}
- Estéril: {prom_esteril}

ESTRUCTURA OBLIGATORIA DEL REPORTE:
1. **RESUMEN EJECUTIVO E ÍTEM 9 (EXPLORACIÓN):** Síntesis metodológica y hallazgos.
2. **FIRMA ADAKÍTICA Y FERTILIDAD MAGMÁTICA:** Discusión sobre relaciones Sr/Y, Y, supresión de plagioclasa y retención de anfíbol/granate.
3. **EVALUACIÓN DE ELEMENTOS TRAZA Y METALOVECTORES:** Análisis comparativo de elementos base (Cu, Zn, Pb) y trazadores metalogénicos (V, Ti, Cr).
4. **ALTERACIÓN GEOQUÍMICA Y TENDENCIAS PETROGENÉTICAS:** Interpretación del enriquecimiento en potasio (Cu vs K) y condiciones redox de la fuente.
5. **CONCLUSIONES Y RECOMENDACIONES TÉCNICAS:** Definición del potencial económico del prospecto y plan de trabajo sugerido.
"""
                                    client = Groq(api_key=groq_api_key)
                                    response = client.chat.completions.create(
                                        model="llama-3.3-70b-versatile",
                                        messages=[
                                            {"role": "system", "content": "Eres un geólogo consultor senior y QP experto en elaboración de informes NI 43-101."},
                                            {"role": "user", "content": prompt_ni43101}
                                        ],
                                        max_tokens=3000,
                                        temperature=0.2,
                                    )
                                    st.session_state.reporte_groq = response.choices[0].message.content
                                except Exception as e:
                                    st.error(f"⚠️ Error al conectar con Groq: {e}")

                    if 'reporte_groq' in st.session_state:
                        st.success("✅ Informe Técnico NI 43-101 Generado Con Éxito.")
                        
                        # Generación directa de buffer para la descarga sin diálogos/popups
                        buffer_word = generar_documento_word(
                            st.session_state.reporte_groq, 
                            st.session_state.get('figuras_informe', {})
                        )

                        st.download_button(
                            label="📥 Descargar Documento con Gráficos Multivariados (.docx)",
                            data=buffer_word,
                            file_name="Informe_Tecnico_NI43101_Multivariado.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                            
                        st.markdown("---")
                        st.markdown(st.session_state.reporte_groq)

            except KeyError as e:
                st.error(f"⚠️ Falta la columna {e} en el archivo subido.")
            except Exception as e:
                st.error(f"⚠️ Error al procesar: {e}")
else:
    st.info("👈 Carga la matriz geoquímica en el panel de control lateral para iniciar.")
